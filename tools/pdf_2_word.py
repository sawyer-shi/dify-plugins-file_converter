import os
import tempfile
from collections.abc import Generator
from typing import Any, Dict, Optional
import json
import time
import io

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from dify_plugin.file.file import File

try:
    import fitz  # PyMuPDF
    from docx import Document
    from docx.shared import Inches
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

PYPDF2_AVAILABLE = PYMUPDF_AVAILABLE  # 兼容性变量

class PdfToWordTool(Tool):
    """Tool for converting PDF documents to Word format."""
    
    def get_file_info(self, file: File) -> dict:
        """
        获取文件信息
        Args:
            file: 文件对象
        Returns:
            文件信息字典
        """
        return {
            "filename": file.filename,
            "extension": file.extension,
            "mime_type": file.mime_type,
            "size": file.size,
            "url": file.url
        }
    
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        try:
            # Get parameters
            file = tool_parameters.get("input_file")
            output_format = tool_parameters.get("output_format", "docx")
            
            if not file:
                yield self.create_text_message("Error: Missing required parameter 'input_file'")
                return
                
            # Get file info
            file_info = self.get_file_info(file)
                
            # Validate input file format
            if not self._validate_input_file(file_info["extension"]):
                yield self.create_text_message("Error: Invalid file format. Only .pdf files are supported")
                return
                
            # Create temporary directory for output
            with tempfile.TemporaryDirectory() as temp_dir:
                # Save uploaded file to temp directory
                input_path = os.path.join(temp_dir, file_info["filename"])
                with open(input_path, 'wb') as f:
                    f.write(file.blob)
                
                # Update file info with the actual path
                file_info["path"] = input_path
                
                # Process conversion
                result = self._process_conversion(file_info, output_format, temp_dir)
                
                if result["success"]:
                    # Create output file info
                    output_files = []
                    for file_path in result["output_files"]:
                        output_files.append({
                            "filename": os.path.basename(file_path),
                            "size": os.path.getsize(file_path),
                            "path": file_path
                        })
                    
                    # Create JSON response
                    json_response = {
                        "success": True,
                        "conversion_type": "pdf_2_word",
                        "input_file": file_info,
                        "output_files": output_files,
                        "message": result["message"]
                    }
                    
                    # Send text message
                    yield self.create_text_message(f"PDF converted to Word successfully: {result['message']}")
                    
                    # Send JSON message
                    yield self.create_json_message(json_response)
                    
                    # Send output files
                    for file_path in result["output_files"]:
                        yield self.create_blob_message(
                            blob=open(file_path, 'rb').read(), 
                            meta={
                                "filename": os.path.basename(file_path),
                                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            }
                        )
                else:
                    # Send error message
                    yield self.create_text_message(f"Conversion failed: {result['message']}")
                    
        except Exception as e:
            yield self.create_text_message(f"Error during conversion: {str(e)}")
    
    def _detect_merged_cells(self, cells, rows, cols):
        """
        检测合并单元格
        算法：
        1. 垂直合并：连续的空单元格 + 上方有内容 = 垂直合并
        2. 水平合并：连续的相同内容 = 水平合并
        
        返回: [(start_row, start_col, end_row, end_col), ...]
        """
        merged_ranges = []
        
        # 创建单元格矩阵
        cell_matrix = [[None for _ in range(cols)] for _ in range(rows)]
        for cell_info in cells:
            r, c = cell_info["row"], cell_info["col"]
            if r < rows and c < cols:
                cell_matrix[r][c] = cell_info["text"]
        
        # 检测垂直合并（同一列连续的空单元格）
        for col in range(cols):
            row = 0
            while row < rows:
                # 找到有内容的单元格
                if cell_matrix[row][col]:
                    # 向下查找连续的空单元格
                    merge_end = row
                    for next_row in range(row + 1, rows):
                        if not cell_matrix[next_row][col] or cell_matrix[next_row][col].strip() == "":
                            merge_end = next_row
                        else:
                            break
                    
                    # 如果有连续空单元格，标记为合并
                    if merge_end > row:
                        merged_ranges.append((row, col, merge_end, col))
                        print(f"Detected vertical merge: ({row},{col}) to ({merge_end},{col})")
                        row = merge_end + 1
                    else:
                        row += 1
                else:
                    row += 1
        
        return merged_ranges
    
    def _create_table_from_structure(self, doc, structure):
        """
        根据PDF表格结构创建Word表格
        先创建结构，再填入数据，检测并应用合并单元格
        
        Args:
            doc: Word Document对象
            structure: 表格结构信息
            
        Returns:
            创建的Word表格对象
        """
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.shared import qn
        
        rows = structure["rows"]
        cols = structure["cols"]
        col_widths = structure.get("col_widths", [])
        cells = structure.get("cells", [])
        
        # 1. 创建空表格
        word_table = doc.add_table(rows=rows, cols=cols)
        word_table.style = 'Table Grid'
        word_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # 2. 设置列宽（使用PDF的实际列宽）
        if col_widths and len(col_widths) == cols:
            for col_idx, col_width in enumerate(col_widths):
                for row in word_table.rows:
                    row.cells[col_idx].width = col_width
        
        # 3. 检测合并单元格
        merged_ranges = self._detect_merged_cells(cells, rows, cols)
        
        # 4. 执行合并操作
        for start_row, start_col, end_row, end_col in merged_ranges:
            try:
                # 合并单元格
                if start_row < rows and end_row < rows and start_col < cols and end_col < cols:
                    start_cell = word_table.cell(start_row, start_col)
                    end_cell = word_table.cell(end_row, end_col)
                    start_cell.merge(end_cell)
                    print(f"Merged cells: ({start_row},{start_col}) to ({end_row},{end_col})")
            except Exception as e:
                print(f"Warning: Failed to merge cells ({start_row},{start_col})-({end_row},{end_col}): {e}")
        
        # 5. 填入数据和应用样式
        for cell_info in cells:
            row_idx = cell_info["row"]
            col_idx = cell_info["col"]
            text = cell_info["text"]
            bg_color = cell_info.get("bg_color")
            
            # 跳过空单元格（可能是被合并的）
            if not text or text.strip() == "":
                continue
            
            if row_idx < rows and col_idx < cols:
                try:
                    cell = word_table.cell(row_idx, col_idx)
                    
                    # 清空并设置内容
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(text)
                    
                    # 设置字体
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    
                    # 设置段落格式
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.line_spacing = 1.15
                    
                    # 判断是否为表头（根据背景色或第一行）
                    is_header = (row_idx == 0 or bg_color is not None)
                    
                    if is_header:
                        run.font.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # 应用背景色
                    if bg_color:
                        try:
                            # pdfplumber的颜色是(r, g, b)元组，值范围0-1
                            r = int(bg_color[0] * 255)
                            g = int(bg_color[1] * 255)
                            b = int(bg_color[2] * 255)
                            bg_hex = f'{r:02X}{g:02X}{b:02X}'
                            
                            shading_elm = OxmlElement('w:shd')
                            shading_elm.set(qn('w:val'), 'clear')
                            shading_elm.set(qn('w:color'), 'auto')
                            shading_elm.set(qn('w:fill'), bg_hex)
                            cell._element.get_or_add_tcPr().append(shading_elm)
                        except Exception as e:
                            print(f"Warning: Failed to apply bg color: {e}")
                    
                    # 设置单元格垂直对齐
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                except Exception as e:
                    print(f"Warning: Failed to format cell ({row_idx},{col_idx}): {e}")
        
        return word_table
    
    def _format_table(self, word_table, table_data, cells_info=None):
        """
        Format the Word table with proper styling and preserve headers.
        严格按照PDF格式，防止文字溢出
        
        Args:
            word_table: python-docx的Table对象
            table_data: 表格数据（二维列表）
            cells_info: PDF中提取的单元格格式信息（背景色等）
        """
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import nsmap
        from docx.oxml import OxmlElement
        from docx.oxml.shared import qn
        
        # 确保所有行的列数一致（处理不规则表格）
        if not table_data:
            return
        
        num_cols = max(len(row) for row in table_data)
        
        # 规范化表格数据：确保所有行都有相同列数
        normalized_data = []
        for row in table_data:
            # 补齐列数
            normalized_row = list(row) + [""] * (num_cols - len(row))
            # 只取前num_cols列
            normalized_data.append(normalized_row[:num_cols])
        
        table_data = normalized_data
        
        # 计算每列的最大内容长度（考虑换行和字符宽度）
        col_max_lengths = [0] * num_cols
        
        for row in table_data:
            for col_idx in range(num_cols):
                if col_idx < len(row) and row[col_idx] is not None:
                    text = str(row[col_idx])
                    # 按换行符分割，取最长的一行
                    lines = text.split('\n')
                    max_line_length = 0
                    for line in lines:
                        # 中文字符宽度系数
                        length = sum(1.8 if ord(c) > 127 else 1 for c in line)
                        max_line_length = max(max_line_length, length)
                    col_max_lengths[col_idx] = max(col_max_lengths[col_idx], max_line_length)
        
        print(f"Column max lengths: {col_max_lengths}")
        
        # 计算列宽（基于内容，使用Cm单位更精确）
        total_length = sum(col_max_lengths)
        available_width_cm = 16.0  # A4纸宽度（21cm）- 左右边距（各2.5cm）
        
        if total_length > 0:
            col_widths = []
            for length in col_max_lengths:
                # 按比例分配宽度
                if length == 0:
                    # 空列，给最小宽度
                    width_cm = 1.5
                else:
                    # 基础宽度 + 按比例分配
                    width_cm = 1.5 + (length / total_length) * (available_width_cm - 1.5 * num_cols)
                    # 限制范围：最小1.5cm，最大5cm
                    width_cm = max(1.5, min(width_cm, 5.0))
                col_widths.append(Cm(width_cm))
            
            print(f"Column widths (cm): {[f'{w.cm:.2f}' for w in col_widths]}")
        else:
            # 平均分配
            col_widths = [Cm(available_width_cm / num_cols)] * num_cols
        
        # Fill the table with data and apply formatting
        for row_idx, row in enumerate(table_data):
            # 设置行高为自动，允许扩展
            try:
                word_table.rows[row_idx].height = None
                word_table.rows[row_idx].height_rule = None  # 自动行高
            except:
                pass
            
            for col_idx, cell_data in enumerate(row):
                cell = word_table.cell(row_idx, col_idx)
                
                # 设置列宽
                if col_idx < len(col_widths):
                    cell.width = col_widths[col_idx]
                
                # 设置单元格边距（减小内边距）
                try:
                    tc = cell._element
                    tcPr = tc.get_or_add_tcPr()
                    tcMar = OxmlElement('w:tcMar')
                    
                    for margin_name in ['top', 'left', 'bottom', 'right']:
                        node = OxmlElement(f'w:{margin_name}')
                        node.set(qn('w:w'), '50')  # 50 twips = 很小的边距
                        node.set(qn('w:type'), 'dxa')
                        tcMar.append(node)
                    
                    tcPr.append(tcMar)
                except Exception as e:
                    print(f"Warning: Failed to set cell margins: {e}")
                
                # 清空默认内容
                cell.text = ""
                
                # 添加内容
                if cell_data is not None and str(cell_data).strip():
                    cell_text = str(cell_data).strip()
                    
                    # 清除默认段落
                    if cell.paragraphs:
                        p = cell.paragraphs[0]
                    else:
                        p = cell.add_paragraph()
                    
                    # 设置段落格式（紧凑）
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.line_spacing = 1.15  # 稍微宽松，避免文字重叠
                    
                    # 添加文字run
                    run = p.add_run(cell_text)
                    
                    # 设置字体（根据内容长度动态调整）
                    cell_length = len(cell_text)
                    if cell_length > 100:
                        font_size = 7  # 内容很多，用更小字体
                    elif cell_length > 50:
                        font_size = 8
                    else:
                        font_size = 9
                    
                    run.font.size = Pt(font_size)
                    
                    # 设置字体名称（确保支持中文）
                    run.font.name = 'Calibri'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    
                    # 检测是否为表头（第一行或有背景色的行）
                    is_header = (row_idx == 0)
                    cell_bg_color = None
                    
                    # 尝试从cells_info中获取该单元格的背景色
                    if cells_info:
                        # 获取当前单元格的bbox（近似）
                        # 注意：这是一个简化实现，精确匹配需要更复杂的算法
                        try:
                            # 根据单元格在表格中的位置估算bbox
                            # 这里简化处理，只检测整行的背景色
                            for rect_info in cells_info:
                                rect_bbox = rect_info["bbox"]
                                fill_color = rect_info.get("fill")
                                
                                # 如果有填充色
                                if fill_color and fill_color != (1, 1, 1):  # 不是白色
                                    cell_bg_color = fill_color
                                    is_header = True  # 有背景色的可能是表头
                                    break
                        except:
                            pass
                    
                    # 应用格式
                    if is_header or row_idx == 0:
                        run.font.bold = True
                        run.font.size = Pt(9)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # 应用背景色
                        try:
                            shading_elm = OxmlElement('w:shd')
                            shading_elm.set(qn('w:val'), 'clear')
                            shading_elm.set(qn('w:color'), 'auto')
                            
                            # 如果有提取的背景色，使用它；否则使用默认浅灰色
                            if cell_bg_color:
                                # pdfplumber的颜色是(r, g, b)元组，值范围0-1
                                r = int(cell_bg_color[0] * 255) if len(cell_bg_color) > 0 else 231
                                g = int(cell_bg_color[1] * 255) if len(cell_bg_color) > 1 else 230
                                b = int(cell_bg_color[2] * 255) if len(cell_bg_color) > 2 else 230
                                bg_hex = f'{r:02X}{g:02X}{b:02X}'
                                shading_elm.set(qn('w:fill'), bg_hex)
                            else:
                                # 默认浅灰色
                                shading_elm.set(qn('w:fill'), 'E7E6E6')
                            
                            cell._element.get_or_add_tcPr().append(shading_elm)
                        except Exception as e:
                            print(f"Warning: Failed to apply cell shading: {e}")
                    else:
                        # 数据行左对齐
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # 设置单元格垂直顶部对齐
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    
                    # 强制单元格宽度（允许换行）
                    try:
                        tcPr = cell._element.get_or_add_tcPr()
                        
                        # 设置单元格宽度
                        tcW = OxmlElement('w:tcW')
                        tcW.set(qn('w:w'), str(int(col_widths[col_idx].twips)))  # 使用twips单位
                        tcW.set(qn('w:type'), 'dxa')
                        tcPr.append(tcW)
                        
                        # 【重要】不添加noWrap，允许文字自动换行
                        # 设置文字方向为从左到右
                        textDirection = OxmlElement('w:textDirection')
                        textDirection.set(qn('w:val'), 'lrTb')  # left-to-right, top-to-bottom
                        tcPr.append(textDirection)
                    except Exception as e:
                        print(f"Warning: Failed to set cell width: {e}")
        
        # Set table alignment to left
        word_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # 禁用自动调整
        word_table.autofit = False
        try:
            word_table.allow_autofit = False
        except:
            pass
        
        # 设置表格样式为简单网格
        try:
            word_table.style = 'Table Grid'
        except:
            pass
    
    def _validate_input_file(self, file_extension: str) -> bool:
        """Validate if the input file format is supported for PDF to Word conversion."""
        return file_extension.lower() == ".pdf"
    
    def _analyze_table_structure(self, table, page):
        """
        分析PDF表格的实际结构
        返回: {
            "rows": 行数,
            "cols": 列数,
            "cells": [{row, col, rowspan, colspan, text, bg_color}, ...]
        }
        """
        structure = {
            "rows": 0,
            "cols": 0,
            "cells": [],
            "col_widths": [],
            "row_heights": []
        }
        
        try:
            # 获取表格的行和列
            rows = table.rows
            cols = table.cols
            
            structure["rows"] = len(rows) - 1 if len(rows) > 0 else 0  # 减1因为包含边界
            structure["cols"] = len(cols) - 1 if len(cols) > 0 else 0
            
            # 计算列宽（基于PDF的实际列坐标）
            if len(cols) > 1:
                col_widths_pts = []
                for i in range(len(cols) - 1):
                    width = cols[i+1] - cols[i]
                    col_widths_pts.append(width)
                
                # 转换为厘米
                from docx.shared import Cm
                # PDF坐标单位是点（pt），1 pt = 0.0353 cm
                structure["col_widths"] = [Cm(w * 0.0353) for w in col_widths_pts]
            
            # 计算行高
            if len(rows) > 1:
                row_heights_pts = []
                for i in range(len(rows) - 1):
                    height = rows[i+1] - rows[i]
                    row_heights_pts.append(height)
                structure["row_heights"] = row_heights_pts
            
            # 提取每个单元格的信息
            bbox = table.bbox
            
            # 获取表格区域内的文本和背景色
            # 使用pdfplumber的chars来获取文本和颜色
            chars = page.chars
            table_chars = [c for c in chars if 
                          c['x0'] >= bbox[0] and c['x1'] <= bbox[2] and
                          c['top'] >= bbox[1] and c['bottom'] <= bbox[3]]
            
            # 获取背景色矩形
            rects = page.rects
            table_rects = [r for r in rects if
                          r['x0'] >= bbox[0] - 2 and r['x1'] <= bbox[2] + 2 and
                          r['top'] >= bbox[1] - 2 and r['bottom'] <= bbox[3] + 2]
            
            # 为每个单元格创建信息
            for row_idx in range(structure["rows"]):
                for col_idx in range(structure["cols"]):
                    # 计算单元格边界
                    cell_x0 = cols[col_idx]
                    cell_x1 = cols[col_idx + 1]
                    cell_y0 = rows[row_idx]
                    cell_y1 = rows[row_idx + 1]
                    
                    # 提取单元格内的文本
                    cell_chars = [c for c in table_chars if
                                 c['x0'] >= cell_x0 and c['x1'] <= cell_x1 and
                                 c['top'] >= cell_y0 and c['bottom'] <= cell_y1]
                    
                    cell_text = ''.join([c['text'] for c in cell_chars])
                    
                    # 检测背景色
                    bg_color = None
                    for rect in table_rects:
                        # 检查矩形是否覆盖该单元格
                        if (rect['x0'] <= cell_x0 + 2 and rect['x1'] >= cell_x1 - 2 and
                            rect['top'] <= cell_y0 + 2 and rect['bottom'] >= cell_y1 - 2):
                            fill = rect.get('non_stroking_color')
                            if fill and fill != (1, 1, 1):  # 不是白色
                                bg_color = fill
                                break
                    
                    structure["cells"].append({
                        "row": row_idx,
                        "col": col_idx,
                        "rowspan": 1,  # 暂时不检测合并
                        "colspan": 1,
                        "text": cell_text.strip(),
                        "bg_color": bg_color
                    })
            
        except Exception as e:
            print(f"Warning: Failed to analyze table structure: {e}")
        
        return structure
    
    def _extract_tables_with_pdfplumber(self, pdf_path, page_num):
        """
        使用pdfplumber提取表格结构和数据
        返回: [{"structure": table_structure, "bbox": bbox}, ...]
        """
        tables_info = []
        
        if not PDFPLUMBER_AVAILABLE:
            return tables_info
        
        try:
            with pdfplumber.open(pdf_path) as pdf:
                if page_num >= len(pdf.pages):
                    return tables_info
                
                page = pdf.pages[page_num]
                
                # 提取所有表格
                tables = page.find_tables(table_settings={
                    "vertical_strategy": "lines",
                    "horizontal_strategy": "lines",
                    "intersection_tolerance": 3,
                    "join_tolerance": 3
                })
                
                for table in tables:
                    try:
                        # 分析表格结构
                        structure = self._analyze_table_structure(table, page)
                        
                        if structure and structure["rows"] > 0 and structure["cols"] > 0:
                            bbox = table.bbox
                            
                            tables_info.append({
                                "structure": structure,
                                "bbox": bbox
                            })
                            
                            print(f"Analyzed table: {structure['rows']} rows x {structure['cols']} cols")
                            print(f"  Column widths: {[f'{w.cm:.2f}cm' for w in structure['col_widths']]}")
                            
                    except Exception as e:
                        print(f"Warning: pdfplumber failed to extract table: {e}")
                        continue
        
        except Exception as e:
            print(f"Warning: pdfplumber processing failed: {e}")
        
        return tables_info
    
    def _get_elements_with_position(self, page, pdf_document, pdf_path=None, page_num=0):
        """
        获取页面所有元素及其位置，按照从上到下的顺序排列
        返回: [(y_position, element_type, element_data), ...]
        
        优先使用pdfplumber提取表格（更准确，支持合并单元格）
        fallback到PyMuPDF的find_tables()
        
        关键改进：先提取表格，再提取文本，排除表格区域避免重复
        """
        elements = []
        
        # 先收集所有表格的区域，用于后续排除
        table_regions = []
        
        # 步骤1：先提取表格（优先pdfplumber）
        # 2. 获取表格及其位置
        # 优先使用pdfplumber（更准确，支持合并单元格）
        tables_extracted = False
        
        if PDFPLUMBER_AVAILABLE and pdf_path:
            try:
                print(f"Using pdfplumber to extract tables on page {page_num + 1}")
                pdfplumber_tables = self._extract_tables_with_pdfplumber(pdf_path, page_num)
                
                if pdfplumber_tables:
                    print(f"pdfplumber found {len(pdfplumber_tables)} tables")
                    for table_info in pdfplumber_tables:
                        bbox = table_info["bbox"]
                        y_position = bbox[1]  # top coordinate
                        structure = table_info["structure"]
                        
                        if structure and structure["rows"] > 0 and structure["cols"] > 0:
                            elements.append((
                                y_position,
                                "table_structured",  # 新类型：结构化表格
                                {
                                    "structure": structure,
                                    "bbox": bbox
                                }
                            ))
                            # 记录表格区域，用于排除文本
                            table_regions.append(bbox)
                    tables_extracted = True
            except Exception as e:
                print(f"Warning: pdfplumber table extraction failed: {e}")
        
        # Fallback到PyMuPDF的find_tables
        if not tables_extracted:
            try:
                print(f"Using PyMuPDF to extract tables on page {page_num + 1}")
                tables = page.find_tables(
                    vertical_strategy="lines",
                    horizontal_strategy="lines",
                    snap_tolerance=3,
                    join_tolerance=3,
                    edge_min_length=3,
                    min_words_vertical=3,
                    min_words_horizontal=1,
                )
                
                if tables.tables:
                    print(f"PyMuPDF found {len(tables.tables)} tables")
                    
                    for table_idx, table in enumerate(tables.tables):
                        try:
                            bbox = table.bbox
                            y_position = bbox[1]
                            
                            # 提取表格数据
                            table_data = table.extract()
                            
                            if table_data and len(table_data) > 0:
                                # 清理表格数据
                                cleaned_data = []
                                for row in table_data:
                                    cleaned_row = []
                                    for cell in row:
                                        if cell is None:
                                            cleaned_row.append("")
                                        else:
                                            cell_text = str(cell).strip()
                                            cleaned_row.append(cell_text)
                                    cleaned_data.append(cleaned_row)
                                
                                # 过滤空行
                                cleaned_data = [row for row in cleaned_data if any(cell for cell in row)]
                                
                                if cleaned_data:
                                    print(f"  Table {table_idx + 1}: {len(cleaned_data)} rows x {len(cleaned_data[0])} cols")
                                    elements.append((
                                        y_position,
                                        "table",
                                        {
                                            "data": cleaned_data,
                                            "bbox": bbox,
                                            "cells_info": None
                                        }
                                    ))
                                    # 记录表格区域
                                    table_regions.append(bbox)
                        except Exception as e:
                            print(f"Warning: Failed to extract table {table_idx}: {e}")
                            continue
            except Exception as e:
                print(f"Warning: PyMuPDF table extraction failed: {e}")
        
        # 步骤2：提取文本块（排除表格区域）
        try:
            text_dict = page.get_text("dict")
            blocks = text_dict.get("blocks", [])
            
            for block in blocks:
                if "lines" in block:
                    # 文本块
                    bbox = block.get("bbox", (0, 0, 0, 0))
                    
                    # 检查文本块是否在表格区域内（避免重复）
                    is_in_table = False
                    for table_bbox in table_regions:
                        # 检查bbox是否重叠
                        # table_bbox格式：(x0, y0, x1, y1) 或 (x0, top, x1, bottom)
                        if (bbox[0] >= table_bbox[0] - 5 and  # x0
                            bbox[2] <= table_bbox[2] + 5 and  # x1
                            bbox[1] >= table_bbox[1] - 5 and  # y0
                            bbox[3] <= table_bbox[3] + 5):    # y1
                            is_in_table = True
                            print(f"Skipping text block in table region: {bbox}")
                            break
                    
                    # 如果文本块在表格内，跳过
                    if is_in_table:
                        continue
                    
                    y_position = bbox[1]  # top y coordinate
                    x_position = bbox[0]  # left x coordinate (用于检测缩进)
                    
                    # 提取文本内容和格式
                    block_text = ""
                    font_size = 12  # 默认字体大小
                    is_bold = False
                    color = None  # RGB颜色
                    
                    # 保存每一行的信息，而不是合并成一个块
                    lines_data = []
                    for line in block["lines"]:
                        line_text = ""
                        line_font_size = 12
                        line_is_bold = False
                        line_color = None
                        
                        for span in line.get("spans", []):
                            line_text += span.get("text", "")
                            # 获取字体信息
                            if "size" in span:
                                line_font_size = span["size"]
                            if "flags" in span:
                                # flags & 16 表示粗体
                                line_is_bold = (span["flags"] & 16) != 0
                            # 获取颜色信息 (RGB格式)
                            if "color" in span:
                                line_color = span["color"]
                        
                        if line_text.strip():
                            lines_data.append({
                                "text": line_text.strip(),
                                "font_size": line_font_size,
                                "is_bold": line_is_bold,
                                "color": line_color
                            })
                        
                        block_text += line_text + "\n"
                        font_size = line_font_size
                        is_bold = line_is_bold
                        color = line_color
                    
                    if block_text.strip():
                        elements.append((
                            y_position,
                            "text",
                            {
                                "text": block_text.strip(),
                                "font_size": font_size,
                                "is_bold": is_bold,
                                "color": color,
                                "x_position": x_position,
                                "bbox": bbox,
                                "lines": lines_data  # 保存每一行的详细信息
                            }
                        ))
        except Exception as e:
            print(f"Warning: Failed to extract text blocks: {e}")
        
        # 步骤3：获取图片及其位置
        try:
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                try:
                    # 获取图片位置
                    xref = img[0]
                    # 获取图片在页面上的位置
                    img_rects = page.get_image_rects(xref)
                    
                    if img_rects:
                        bbox = img_rects[0]
                        y_position = bbox[1]  # top y coordinate
                        
                        # 获取图片数据
                        pix = fitz.Pixmap(pdf_document, xref)
                        
                        # Skip CMYK images
                        if pix.n - pix.alpha < 4:
                            img_data = pix.tobytes("png")
                            
                            elements.append((
                                y_position,
                                "image",
                                {
                                    "data": img_data,
                                    "bbox": bbox,
                                    "width": pix.width,
                                    "height": pix.height
                                }
                            ))
                        
                        pix = None
                except Exception as e:
                    print(f"Warning: Failed to extract image {img_index}: {e}")
                    continue
        except Exception as e:
            print(f"Warning: Failed to get images: {e}")
        
        # 按照y坐标排序（从上到下）
        elements.sort(key=lambda x: x[0])
        
        return elements
    
    def _process_conversion(self, file_info: Dict[str, Any], output_format: str, temp_dir: str) -> Dict[str, Any]:
        """
        Process the PDF to Word conversion using PyMuPDF library.
        关键改进：按照PDF的实际布局顺序（从上到下）处理内容，不添加额外内容
        """
        input_path = file_info["path"]
        output_files = []
        
        try:
            if not PYPDF2_AVAILABLE:
                return {"success": False, "message": "Required libraries (PyMuPDF, python-docx) are not available for PDF conversion"}
            
            # Default to docx if not specified
            if not output_format:
                output_format = "docx"
            elif output_format.lower() not in ["docx"]:
                output_format = "docx"
            
            # Generate output file path
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(temp_dir, f"{base_name}.{output_format}")
            
            # Open the PDF file
            try:
                pdf_document = fitz.open(input_path)
            except Exception as e:
                return {"success": False, "message": f"Failed to open PDF file: {str(e)}"}
            
            # Create a new Word document
            doc = Document()
            
            # Process each page
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                
                # 获取页面所有元素并按位置排序（传入pdf_path用于pdfplumber）
                elements = self._get_elements_with_position(page, pdf_document, input_path, page_num)
                
                # 按顺序处理每个元素
                for y_pos, element_type, element_data in elements:
                    if element_type == "text":
                        # 添加文本段落
                        from docx.shared import Pt, RGBColor, Inches
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        
                        text = element_data["text"]
                        font_size = element_data["font_size"]
                        is_bold = element_data["is_bold"]
                        color = element_data.get("color")
                        x_position = element_data.get("x_position", 0)
                        lines = element_data.get("lines", [])
                        
                        # 判断是否为标题（根据字体大小）
                        if font_size >= 16:
                            # 大字体，作为一级标题
                            heading = doc.add_heading(text, level=1)
                            # 应用颜色到标题
                            if color is not None:
                                for run in heading.runs:
                                    # PDF颜色是整数，需要转换为RGB
                                    r = (color >> 16) & 0xFF
                                    g = (color >> 8) & 0xFF
                                    b = color & 0xFF
                                    run.font.color.rgb = RGBColor(r, g, b)
                        elif font_size >= 14:
                            # 中等字体，作为二级标题
                            heading = doc.add_heading(text, level=2)
                            # 应用颜色到标题
                            if color is not None:
                                for run in heading.runs:
                                    r = (color >> 16) & 0xFF
                                    g = (color >> 8) & 0xFF
                                    b = color & 0xFF
                                    run.font.color.rgb = RGBColor(r, g, b)
                        else:
                            # 普通文本 - 逐行处理以保留换行
                            if lines and len(lines) > 1:
                                # 有多行，逐行添加段落保留换行
                                for line_data in lines:
                                    p = doc.add_paragraph()
                                    
                                    # 设置缩进
                                    if x_position > 80:
                                        indent_inches = (x_position - 72) / 72.0
                                        p.paragraph_format.left_indent = Inches(min(indent_inches, 2.0))
                                    
                                    # 添加文本run
                                    run = p.add_run(line_data["text"])
                                    
                                    # 应用格式
                                    if line_data.get("is_bold"):
                                        run.bold = True
                                    
                                    # 应用颜色
                                    line_color = line_data.get("color")
                                    if line_color is not None:
                                        r = (line_color >> 16) & 0xFF
                                        g = (line_color >> 8) & 0xFF
                                        b = line_color & 0xFF
                                        run.font.color.rgb = RGBColor(r, g, b)
                                    
                                    # 设置字体大小
                                    line_font_size = line_data.get("font_size", font_size)
                                    run.font.size = Pt(max(line_font_size, 9))
                            else:
                                # 单行文本，直接添加
                                p = doc.add_paragraph(text)
                                
                                # 设置缩进
                                if x_position > 80:
                                    indent_inches = (x_position - 72) / 72.0
                                    p.paragraph_format.left_indent = Inches(min(indent_inches, 2.0))
                                
                                # 应用格式到run
                                for run in p.runs:
                                    if is_bold:
                                        run.bold = True
                                    # 应用颜色
                                    if color is not None:
                                        r = (color >> 16) & 0xFF
                                        g = (color >> 8) & 0xFF
                                        b = color & 0xFF
                                        run.font.color.rgb = RGBColor(r, g, b)
                                    # 设置字体大小
                                    run.font.size = Pt(max(font_size, 9))
                    
                    elif element_type == "table_structured":
                        # 使用结构化方法创建表格（pdfplumber提取，包含实际列宽和样式）
                        structure = element_data["structure"]
                        self._create_table_from_structure(doc, structure)
                    
                    elif element_type == "table":
                        # 使用传统方法创建表格（PyMuPDF提取的fallback）
                        table_data = element_data["data"]
                        cells_info = element_data.get("cells_info", None)
                        
                        if table_data and len(table_data) > 0:
                            cols = len(table_data[0])
                            word_table = doc.add_table(rows=len(table_data), cols=cols)
                            word_table.style = 'Table Grid'
                            
                            # 填充表格数据（传入cells_info以应用PDF样式）
                            self._format_table(word_table, table_data, cells_info)
                    
                    elif element_type == "image":
                        # 添加图片
                        from docx.shared import Inches as DocxInches
                        
                        img_data = element_data["data"]
                        img_stream = io.BytesIO(img_data)
                        
                        # 根据原始尺寸计算合适的显示宽度
                        width = element_data["width"]
                        height = element_data["height"]
                        
                        # 限制最大宽度为6英寸
                        max_width = 6.0
                        if width > height:
                            doc_width = min(max_width, width / 100.0)  # 简单的缩放
                        else:
                            doc_width = min(4.0, width / 100.0)
                        
                        doc.add_picture(img_stream, width=DocxInches(doc_width))
                
                # 在页面之间添加分页符（除了最后一页）
                if page_num < len(pdf_document) - 1:
                    doc.add_page_break()
            
            # Close the PDF document
            pdf_document.close()
            
            # Save the Word document
            try:
                doc.save(output_path)
            except Exception as e:
                return {"success": False, "message": f"Failed to save Word document: {str(e)}"}
            
            # Check if file exists and has content
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                output_files.append(output_path)
            else:
                return {"success": False, "message": "Output file was not created or is empty"}
            
            return {
                "success": True, 
                "message": f"PDF converted to Word ({output_format}) successfully - preserving original layout order",
                "output_files": output_files
            }
                
        except Exception as e:
            return {"success": False, "message": f"Conversion error: {str(e)}"}