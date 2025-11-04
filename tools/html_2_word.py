import os
import tempfile
import uuid
from typing import Any, Dict, List, Optional, Generator
import shutil
from collections.abc import Generator
from bs4 import BeautifulSoup

# python-docx相关导入
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# docxtpl相关导入
try:
    from docxtpl import DocxTemplate, RichText
    DOCTXTPL_AVAILABLE = True
except ImportError:
    DOCTXTPL_AVAILABLE = False

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from dify_plugin.file.file import File


class HtmlToWordTool(Tool):
    """
    HTML to Word conversion tool using python-docx
    """
    
    def get_runtime_parameters(self) -> List[Dict[str, Any]]:
        return [
            {
                "name": "input_file",
                "type": "file",
                "required": True,
                "description": "要转换的HTML文件"
            },
            {
                "name": "output_file",
                "type": "string",
                "required": False,
                "description": "输出Word文件名（可选，默认使用原文件名.docx）"
            },
            {
                "name": "font_family",
                "type": "string",
                "required": False,
                "description": "字体名称（默认：微软雅黑）",
                "default": "微软雅黑"
            },
            {
                "name": "font_size",
                "type": "number",
                "required": False,
                "description": "字体大小（默认：11）",
                "default": 11
            },
            {
                "name": "encoding",
                "type": "string",
                "required": False,
                "description": "HTML文件编码（默认utf-8）",
                "default": "utf-8"
            }
        ]

    def _invoke(self, tool_parameters: Dict[str, Any]) -> Generator[ToolInvokeMessage]:
        """
        Convert HTML to Word using python-docx
        """
        # 检查依赖
        if not PYTHON_DOCX_AVAILABLE:
            yield self.create_text_message("Error: python-docx库未安装，请安装python-docx>=0.8.11")
            return
        
        # 获取参数
        input_file = tool_parameters.get("input_file")
        output_file = tool_parameters.get("output_file")
        font_family = tool_parameters.get("font_family", "微软雅黑")
        font_size = tool_parameters.get("font_size", 11)
        encoding = tool_parameters.get("encoding", "utf-8")
        
        # 验证输入文件
        if not input_file:
            yield self.create_text_message("Error: 请提供有效的HTML文件")
            return
        
        try:
            # 读取HTML内容
            if hasattr(input_file, 'blob'):
                # 如果是File对象，使用blob属性
                html_content = input_file.blob
            elif hasattr(input_file, 'read'):
                # 如果是文件对象，使用read方法
                html_content = input_file.read()
            else:
                # 如果是字符串路径，直接读取文件
                with open(input_file, 'rb') as f:
                    html_content = f.read()
            
            # 如果是二进制内容，尝试解码
            if isinstance(html_content, bytes):
                try:
                    html_content = html_content.decode(encoding)
                except UnicodeDecodeError:
                    html_content = html_content.decode('utf-8', errors='replace')
            
            # 使用BeautifulSoup解析HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 生成输出文件名
            if not output_file:
                if hasattr(input_file, 'filename'):
                    original_name = input_file.filename
                elif hasattr(input_file, 'name'):
                    original_name = input_file.name
                else:
                    original_name = 'html_document'
                base_name = os.path.splitext(original_name)[0]
                output_file = f"{base_name}.docx"
            
            # 创建临时目录
            with tempfile.TemporaryDirectory() as temp_dir:
                # 设置输出文件完整路径
                output_path = os.path.join(temp_dir, output_file)
                
                # 创建Word文档
                doc = Document()
                
                # 设置默认字体
                style = doc.styles['Normal']
                font = style.font
                font.name = font_family
                font.size = Pt(font_size)
                
                # 处理HTML内容
                self._process_html_elements(soup, doc, font_family, font_size)
                
                # 保存文档
                doc.save(output_path)
                
                # 获取文件大小
                file_size = os.path.getsize(output_path)
                file_size_mb = round(file_size / (1024 * 1024), 2)
                
                # 提取文本内容用于预览
                text_content = soup.get_text()
                preview_text = text_content[:500] + "..." if len(text_content) > 500 else text_content
                
                # 创建JSON响应
                json_response = {
                    "success": True,
                    "conversion_type": "html_2_word",
                    "input_file": {
                        "filename": getattr(input_file, 'filename', getattr(input_file, 'name', 'html_document')),
                        "size": len(html_content) if isinstance(html_content, bytes) else len(html_content.encode())
                    },
                    "output_file": {
                        "filename": output_file,
                        "size": file_size_mb,
                        "path": output_path
                    },
                    "preview_text": preview_text,
                    "message": f"HTML文件已成功转换为Word格式\n输出文件: {output_file}\n文件大小: {file_size_mb} MB\n字体: {font_family}, 字号: {font_size}"
                }
                
                # 发送文本消息
                yield self.create_text_message(f"HTML文件已成功转换为Word格式: {output_file}")
                
                # 发送JSON消息
                yield self.create_json_message(json_response)
                
                # 发送文件
                with open(output_path, 'rb') as f:
                    yield self.create_blob_message(
                        blob=f.read(), 
                        meta={
                            "filename": output_file,
                            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        }
                    )
                
        except Exception as e:
            yield self.create_text_message(f"转换过程中发生错误: {str(e)}")
    
    def _process_html_elements(self, soup, doc, font_family, font_size):
        """处理HTML元素并转换为Word格式"""
        # 处理标题
        for i in range(1, 7):
            headings = soup.find_all(f'h{i}')
            for heading in headings:
                text = heading.get_text().strip()
                if text:
                    # 添加标题
                    heading_para = doc.add_heading(text, level=i)
                    # 设置字体
                    for run in heading_para.runs:
                        run.font.name = font_family
                        run.font.size = Pt(font_size + (6 - i))  # 标题字号稍大
        
        # 处理段落
        paragraphs = soup.find_all('p')
        for p in paragraphs:
            text = p.get_text().strip()
            if text:
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.font.name = font_family
                run.font.size = Pt(font_size)
        
        # 处理其他文本内容（不在标签内的文本）
        if not soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p']):
            # 如果没有找到任何标题或段落，则提取所有文本
            text = soup.get_text().strip()
            if text:
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.font.name = font_family
                run.font.size = Pt(font_size)
        
        # 处理列表
        lists = soup.find_all(['ul', 'ol'])
        for lst in lists:
            items = lst.find_all('li', recursive=False)
            for item in items:
                text = item.get_text().strip()
                if text:
                    if lst.name == 'ul':
                        para = doc.add_paragraph(text, style='List Bullet')
                    else:
                        para = doc.add_paragraph(text, style='List Number')
                    
                    # 设置字体
                    for run in para.runs:
                        run.font.name = font_family
                        run.font.size = Pt(font_size)
        
        # 处理表格
        tables = soup.find_all('table')
        for table in tables:
            rows = table.find_all('tr')
            if not rows:
                continue
                
            # 创建Word表格
            # 计算列数
            cols = max(len(row.find_all(['th', 'td'])) for row in rows)
            if cols == 0:
                continue
                
            word_table = doc.add_table(rows=len(rows), cols=cols)
            word_table.style = 'Table Grid'
            
            # 填充表格数据
            for i, row in enumerate(rows):
                cells = row.find_all(['th', 'td'])
                for j, cell in enumerate(cells):
                    if j >= cols:
                        break
                        
                    text = cell.get_text().strip()
                    word_table.cell(i, j).text = text
                    
                    # 设置单元格字体
                    for paragraph in word_table.cell(i, j).paragraphs:
                        for run in paragraph.runs:
                            run.font.name = font_family
                            run.font.size = Pt(font_size)
        
        # 处理图片
        images = soup.find_all('img')
        for img in images:
            src = img.get('src', '')
            alt = img.get('alt', '')
            
            # 如果是本地图片路径
            if src.startswith(('http://', 'https://')):
                # 跳过网络图片，因为需要额外处理
                para = doc.add_paragraph(f"[图片: {alt or src}]")
                para.runs[0].font.name = font_family
                para.runs[0].font.size = Pt(font_size)
            elif os.path.exists(src):
                try:
                    # 添加本地图片
                    width = Inches(6)  # 设置图片宽度
                    doc.add_picture(src, width=width)
                except Exception:
                    # 如果添加图片失败，添加文本描述
                    para = doc.add_paragraph(f"[图片: {alt or src}]")
                    para.runs[0].font.name = font_family
                    para.runs[0].font.size = Pt(font_size)
            else:
                # 添加图片占位符
                para = doc.add_paragraph(f"[图片: {alt or src}]")
                para.runs[0].font.name = font_family
                para.runs[0].font.size = Pt(font_size)
        
        # 处理链接
        links = soup.find_all('a')
        for link in links:
            href = link.get('href', '')
            text = link.get_text().strip()
            if text:
                para = doc.add_paragraph()
                run = para.add_run(f"{text} ({href})")
                run.font.name = font_family
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色