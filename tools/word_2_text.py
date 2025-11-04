import os
import tempfile
import time
from collections.abc import Generator
from typing import Any, Dict, Optional
import json

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from dify_plugin.file.file import File

# Try to import python-docx for Word processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class WordToTextTool(Tool):
    """Tool for converting Word documents to text format."""
    
    def get_file_info(self, file: File) -> dict:
        """
        获取文件信息
        Args:
            file: 文件对象
        Returns:
            文件信息字典
        """
        file_info = {
            "filename": file.filename,
            "extension": file.extension,
            "mime_type": file.mime_type,
            "size": file.size,
            "url": file.url
        }
        
        # Add path attribute if it exists (for MockFile in testing)
        if hasattr(file, 'path'):
            file_info["path"] = file.path
            
        return file_info
    
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        try:
            # Get input file parameter
            file = tool_parameters.get("input_file")
            
            if not file:
                yield self.create_text_message("Error: Missing required parameter 'input_file'")
                return
                
            # Get file info
            file_info = self.get_file_info(file)
                
            # Validate input file format
            if not self._validate_input_file(file_info):
                yield self.create_text_message("Error: Invalid file format. Only .doc and .docx files are supported")
                return
                
            # Create temporary directory for input file
            with tempfile.TemporaryDirectory() as temp_dir:
                # Use the test directory as output directory
                import os
                output_dir = r"D:\Work\Cursor\file_converter\test"
                os.makedirs(output_dir, exist_ok=True)
                
                # Save uploaded file to temp directory
                input_path = os.path.join(temp_dir, file_info["filename"])
                with open(input_path, 'wb') as f:
                    f.write(file.blob)
                
                # Update file info with the actual path
                file_info["path"] = input_path
                
                # Process conversion
                result = self._process_conversion(input_path, temp_dir)
                
                if result["success"]:
                    # Create output file info
                    output_files = []
                    for output_file_info in result["output_files"]:
                        output_files.append({
                            "filename": output_file_info["filename"],
                            "size": len(output_file_info["content"]),
                            "path": output_file_info["path"]
                        })
                    
                    # Create JSON response
                    json_response = {
                        "success": True,
                        "conversion_type": "word_2_text",
                        "input_file": file_info,
                        "output_files": output_files,
                        "message": result["message"]
                    }
                    
                    # Send text message
                    yield self.create_text_message(f"Word document converted to text successfully: {result['message']}")
                    
                    # Send JSON message
                    yield self.create_json_message(json_response)
                    
                    # Send output files
                    for file_info in result["output_files"]:
                        try:
                            # Use the pre-read content
                            if "content" in file_info:
                                yield self.create_blob_message(
                                    blob=file_info["content"], 
                                    meta={
                                        "filename": file_info["filename"],
                                        "mime_type": "text/plain"
                                    }
                                )
                            else:
                                yield self.create_text_message(f"Error: No content available for file {file_info.get('filename', 'unknown')}")
                        except Exception as e:
                            yield self.create_text_message(f"Error sending file: {str(e)}")
                    
                    # Clean up only the temporary directory, not the output directory
                    try:
                        import shutil
                        # Only clean up the temp directory, not the output directory
                        # The temp directory will be automatically cleaned up by the context manager
                        pass
                    except Exception as e:
                        # Ignore cleanup errors
                        pass
                else:
                    # Send error message
                    yield self.create_text_message(f"Conversion failed: {result['message']}")
                    
        except Exception as e:
            yield self.create_text_message(f"Error during conversion: {str(e)}")
    
    def _validate_input_file(self, file_info: dict) -> bool:
        """Validate if the input file is a valid Word document."""
        # Check file extension
        if not file_info["extension"].lower().endswith(('.docx', '.doc')):
            return False
            
        # Check if file is readable by python-docx
        if DOCX_AVAILABLE and "path" in file_info:
            try:
                doc = Document(file_info["path"])
                # Try to access the document to ensure it's valid
                _ = doc.paragraphs[0].text if doc.paragraphs else ""
                return True
            except Exception:
                return False
        
        # If python-docx is not available or path not available, just check file extension
        return True
    
    def _process_conversion(self, input_path: str, temp_dir: str) -> Dict[str, Any]:
        """Process the Word to text conversion using python-docx."""
        output_files = []
        
        # Generate output file path
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        output_path = os.path.join(temp_dir, f"{base_name}.txt")
        
        # Check if required libraries are available
        if not DOCX_AVAILABLE:
            return {"success": False, "message": "Required library (python-docx) is not available. Please install it using: pip install python-docx"}
        
        try:
            # Load the Word document
            doc = Document(input_path)
            
            # Extract text from paragraphs
            text_content = []
            
            # Process paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_content.append(text)
            
            # Process tables
            for table in doc.tables:
                text_content.append("\n--- Table ---")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
                text_content.append("--- End of Table ---\n")
            
            # Join all text content
            full_text = "\n\n".join(text_content)
            
            # Write text to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_text)
            
            # Wait for file to be fully written
            time.sleep(1)
            
            # Check if file exists and has content
            if not os.path.exists(output_path):
                return {"success": False, "message": "Output text file was not created"}
                
            if os.path.getsize(output_path) == 0:
                return {"success": False, "message": "Output text file is empty"}
            
            # Try multiple times to read the file
            file_content = None
            for attempt in range(3):
                try:
                    with open(output_path, 'r', encoding='utf-8') as f:
                        file_content = f.read()
                    break
                except Exception as e:
                    if attempt < 2:
                        time.sleep(1)
                    else:
                        return {"success": False, "message": f"Error reading converted file: {str(e)}"}
            
            if file_content:
                output_files.append({
                    "path": output_path,
                    "content": file_content.encode('utf-8'),
                    "filename": f"{base_name}.txt"
                })
                return {
                    "success": True, 
                    "message": "Word document converted to text successfully",
                    "output_files": output_files
                }
            else:
                return {"success": False, "message": "Failed to read converted file after multiple attempts"}
                    
        except Exception as e:
            return {"success": False, "message": f"Error converting Word to text: {str(e)}"}